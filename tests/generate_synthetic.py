"""E1: Synthetic Test File Generator.

Generates test files with known PII data for deterministic extraction
validation.  Each file has an accompanying manifest (JSON) listing
exactly what PII should be found.

Formats generated:
  - CSV (tabular, single-page)
  - XLSX (multi-sheet, header variations)
  - XLS (legacy, smart header test)
  - DOCX (prose with embedded PII)
  - HTML (structured + unstructured PII)
  - EML (email body + attachment)
  - TXT (plain text with PII patterns)

Usage:
    python -m tests.generate_synthetic [--output-dir tests/fixtures]
"""

from __future__ import annotations

import csv
import json
import os
import sys
from dataclasses import dataclass, asdict
from pathlib import Path

# ---------------------------------------------------------------------------
# Known PII records — ground truth
# ---------------------------------------------------------------------------

@dataclass
class SyntheticPerson:
    name: str
    ssn: str
    dob: str
    email: str
    phone: str
    address: str


PERSONS = [
    SyntheticPerson("John Michael Smith", "123-45-6789", "01/15/1990",
                    "john.smith@example.com", "(555) 123-4567", "123 Main St, Springfield IL 62701"),
    SyntheticPerson("Maria Elena Garcia", "234-56-7890", "03/22/1985",
                    "maria.garcia@example.com", "(555) 234-5678", "456 Oak Ave, Portland OR 97201"),
    SyntheticPerson("Robert James Wilson", "345-67-8901", "07/04/1978",
                    "r.wilson@example.com", "(555) 345-6789", "789 Pine Rd, Austin TX 78701"),
    SyntheticPerson("Sarah Ann Johnson", "456-78-9012", "11/30/1992",
                    "s.johnson@example.com", "(555) 456-7890", "321 Elm Blvd, Denver CO 80201"),
    SyntheticPerson("David Lee Chen", "567-89-0123", "05/18/1988",
                    "david.chen@example.com", "(555) 567-8901", "654 Maple Ln, Seattle WA 98101"),
    SyntheticPerson("Jennifer Rose Brown", "678-90-1234", "09/12/1995",
                    "j.brown@example.com", "(555) 678-9012", "987 Cedar Dr, Miami FL 33101"),
    SyntheticPerson("Michael Anthony Davis", "789-01-2345", "02/28/1982",
                    "m.davis@example.com", "(555) 789-0123", "147 Birch Way, Chicago IL 60601"),
    SyntheticPerson("Emily Kate Martinez", "890-12-3456", "06/15/1993",
                    "e.martinez@example.com", "(555) 890-1234", "258 Walnut Ct, Phoenix AZ 85001"),
]

# Organizational metadata (should NOT be extracted as breach subjects)
ORG_METADATA = {
    "company_name": "Acme Healthcare Systems Inc.",
    "company_phone": "(800) 555-0100",
    "company_email": "privacy@acmehealthcare.com",
    "sender_name": "Dr. Patricia Williams",
    "sender_title": "Chief Privacy Officer",
}


def _manifest(persons: list[SyntheticPerson], filename: str, fmt: str, notes: str = "") -> dict:
    """Build a manifest dict for ground truth."""
    return {
        "filename": filename,
        "format": fmt,
        "expected_records": len(persons),
        "notes": notes,
        "persons": [asdict(p) for p in persons],
        "org_metadata": ORG_METADATA,
        "false_positives_expected": [
            ORG_METADATA["company_name"],
            ORG_METADATA["company_phone"],
            ORG_METADATA["sender_name"],
        ],
    }


# ---------------------------------------------------------------------------
# File generators
# ---------------------------------------------------------------------------

def generate_csv(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate a CSV with tabular PII data."""
    fname = "synthetic_pii.csv"
    path = output_dir / fname

    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Full Name", "SSN", "Date of Birth", "Email", "Phone", "Address"])
        for p in persons:
            writer.writerow([p.name, p.ssn, p.dob, p.email, p.phone, p.address])

    return _manifest(persons, fname, "csv", "Simple tabular CSV with header row")


def generate_xlsx(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate an XLSX with multiple sheets and header variations."""
    fname = "synthetic_pii.xlsx"
    path = output_dir / fname

    try:
        import openpyxl
    except ImportError:
        print("SKIP: openpyxl not installed, skipping XLSX generation")
        return _manifest(persons[:4], fname, "xlsx", "SKIPPED — openpyxl not available")

    wb = openpyxl.Workbook()

    # Sheet 1: Standard headers
    ws1 = wb.active
    ws1.title = "Breach Subjects"
    ws1.append(["Full Name", "SSN", "Date of Birth", "Email", "Phone", "Address"])
    for p in persons[:4]:
        ws1.append([p.name, p.ssn, p.dob, p.email, p.phone, p.address])

    # Sheet 2: Title row + blank row + headers (tests D3 smart header detection)
    ws2 = wb.create_sheet("Report Summary")
    ws2.append(["Acme Healthcare — Breach Report 2026"])
    ws2.append([])  # blank row
    ws2.append(["Name", "Government ID", "DOB", "Contact Email", "Contact Phone", "Home Address"])
    for p in persons[4:]:
        ws2.append([p.name, p.ssn, p.dob, p.email, p.phone, p.address])

    # Sheet 3: Hidden sheet (should be skipped)
    ws3 = wb.create_sheet("Internal Notes")
    ws3.sheet_state = "hidden"
    ws3.append(["This sheet should NOT be processed"])

    wb.save(str(path))
    return _manifest(persons, fname, "xlsx",
                     "Multi-sheet: standard headers + title-row offset + hidden sheet")


def generate_xls(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate a legacy XLS with a title row before headers (D3 test)."""
    fname = "synthetic_pii_legacy.xls"
    path = output_dir / fname

    try:
        import xlwt
    except ImportError:
        # Create a minimal CSV as fallback
        print("SKIP: xlwt not installed, creating CSV fallback for XLS test")
        fallback = output_dir / fname.replace(".xls", "_fallback.csv")
        with open(fallback, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["Breach Notification Report"])
            writer.writerow([])
            writer.writerow(["Name", "SSN", "DOB", "Email"])
            for p in persons[:4]:
                writer.writerow([p.name, p.ssn, p.dob, p.email])
        return _manifest(persons[:4], fname, "xls", "SKIPPED — xlwt not available, CSV fallback")

    wb = xlwt.Workbook()
    ws = wb.add_sheet("Breach Data")

    # Row 0: Title (not header)
    ws.write(0, 0, "CONFIDENTIAL — Breach Notification List")
    # Row 1: blank
    # Row 2: actual headers
    for col, header in enumerate(["Full Name", "SSN", "DOB", "Email", "Phone"]):
        ws.write(2, col, header)
    # Row 3+: data
    for row_idx, p in enumerate(persons[:4], 3):
        ws.write(row_idx, 0, p.name)
        ws.write(row_idx, 1, p.ssn)
        ws.write(row_idx, 2, p.dob)
        ws.write(row_idx, 3, p.email)
        ws.write(row_idx, 4, p.phone)

    wb.save(str(path))
    return _manifest(persons[:4], fname, "xls",
                     "Legacy XLS with title row 0, blank row 1, headers row 2")


def generate_docx(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate a DOCX with prose containing embedded PII."""
    fname = "synthetic_breach_report.docx"
    path = output_dir / fname

    try:
        from docx import Document
    except ImportError:
        print("SKIP: python-docx not installed, skipping DOCX generation")
        return _manifest(persons[:3], fname, "docx", "SKIPPED — python-docx not available")

    doc = Document()
    doc.add_heading("Breach Notification Report", 0)
    doc.add_paragraph(
        f"Prepared by: {ORG_METADATA['sender_name']}, {ORG_METADATA['sender_title']}\n"
        f"Organization: {ORG_METADATA['company_name']}\n"
        f"Contact: {ORG_METADATA['company_email']} | {ORG_METADATA['company_phone']}\n"
    )
    doc.add_heading("Affected Individuals", level=1)

    for p in persons[:3]:
        doc.add_paragraph(
            f"Name: {p.name}\n"
            f"Social Security Number: {p.ssn}\n"
            f"Date of Birth: {p.dob}\n"
            f"Email: {p.email}\n"
            f"Phone: {p.phone}\n"
            f"Address: {p.address}\n"
        )
        doc.add_paragraph("---")

    doc.save(str(path))
    return _manifest(persons[:3], fname, "docx",
                     "Prose doc with org metadata header and embedded PII")


def generate_html(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate an HTML file with a PII table."""
    fname = "synthetic_breach_table.html"
    path = output_dir / fname

    rows_html = ""
    for p in persons[:5]:
        rows_html += f"""        <tr>
            <td>{p.name}</td><td>{p.ssn}</td><td>{p.dob}</td>
            <td>{p.email}</td><td>{p.phone}</td><td>{p.address}</td>
        </tr>\n"""

    html = f"""<!DOCTYPE html>
<html>
<head><title>Breach Notification — {ORG_METADATA['company_name']}</title></head>
<body>
    <h1>Affected Individuals Report</h1>
    <p>Prepared by {ORG_METADATA['sender_name']}, {ORG_METADATA['company_name']}</p>
    <p>Contact: {ORG_METADATA['company_email']} | {ORG_METADATA['company_phone']}</p>
    <table border="1">
        <tr><th>Name</th><th>SSN</th><th>DOB</th><th>Email</th><th>Phone</th><th>Address</th></tr>
{rows_html}    </table>
</body>
</html>"""

    path.write_text(html)
    return _manifest(persons[:5], fname, "html", "HTML table with org metadata header")


def generate_eml(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate an EML file with PII in body (D2 attachment test)."""
    fname = "synthetic_breach_email.eml"
    path = output_dir / fname

    body_lines = [
        f"The following individuals were affected by the data breach:\n",
    ]
    for p in persons[:3]:
        body_lines.append(f"- {p.name}, SSN: {p.ssn}, DOB: {p.dob}, Email: {p.email}")

    body_lines.append(f"\nPlease contact {ORG_METADATA['sender_name']} at {ORG_METADATA['company_email']}")
    body = "\n".join(body_lines)

    eml = f"""From: {ORG_METADATA['sender_name']} <{ORG_METADATA['company_email']}>
To: legal@acmehealthcare.com
Subject: CONFIDENTIAL — Breach Notification List
Date: Mon, 6 Apr 2026 10:00:00 -0500
Content-Type: text/plain; charset="utf-8"

{body}
"""

    path.write_text(eml)
    return _manifest(persons[:3], fname, "eml",
                     "Plain-text email with PII in body, sender in From header")


def generate_txt(output_dir: Path, persons: list[SyntheticPerson]) -> dict:
    """Generate a plain text file with PII patterns."""
    fname = "synthetic_pii_list.txt"
    path = output_dir / fname

    lines = [
        f"BREACH NOTIFICATION LIST — {ORG_METADATA['company_name']}",
        f"Prepared by: {ORG_METADATA['sender_name']}",
        f"Contact: {ORG_METADATA['company_phone']}",
        "",
        "=" * 60,
        "",
    ]

    for p in persons:
        lines.extend([
            f"Name: {p.name}",
            f"SSN: {p.ssn}",
            f"DOB: {p.dob}",
            f"Email: {p.email}",
            f"Phone: {p.phone}",
            f"Address: {p.address}",
            "-" * 40,
        ])

    path.write_text("\n".join(lines))
    return _manifest(persons, fname, "txt", "Plain text with all 8 persons, org metadata header")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate_all(output_dir: str = "tests/fixtures") -> list[dict]:
    """Generate all synthetic test files and return manifests."""
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    manifests = []
    manifests.append(generate_csv(out, PERSONS))
    manifests.append(generate_xlsx(out, PERSONS))
    manifests.append(generate_xls(out, PERSONS))
    manifests.append(generate_docx(out, PERSONS))
    manifests.append(generate_html(out, PERSONS))
    manifests.append(generate_eml(out, PERSONS))
    manifests.append(generate_txt(out, PERSONS))

    # Write combined manifest
    manifest_path = out / "manifest.json"
    with open(manifest_path, "w") as f:
        json.dump(manifests, f, indent=2)

    print(f"Generated {len(manifests)} test files in {out}/")
    print(f"Manifest: {manifest_path}")
    return manifests


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "tests/fixtures"
    generate_all(output)
