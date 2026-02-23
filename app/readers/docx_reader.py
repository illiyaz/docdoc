"""DOCX reader: python-docx paragraph and table extraction.

Emits prose blocks for paragraphs and table_cell / table_header blocks
for table content. Each document is treated as a single logical unit;
page_or_sheet is set to 0 (DOCX has no reliable page boundaries at
parse time).

Table rows use a shared table_id UUID. The first row of each table is
emitted as block_type="table_header" with col_header populated from
the cell text.
"""
from __future__ import annotations

import uuid
from pathlib import Path

import docx

from app.readers.base import BaseReader, ExtractedBlock


class DOCXReader(BaseReader):
    """Extract content from a .docx file and emit ExtractedBlock objects."""

    def __init__(self, path: str | Path) -> None:
        super().__init__(path)

    def read(self) -> list[ExtractedBlock]:
        """Extract paragraphs and tables; return ExtractedBlock list."""
        source = str(self.path)
        doc = docx.Document(str(self.path))
        all_blocks: list[ExtractedBlock] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_blocks.append(ExtractedBlock(
                    text=text,
                    page_or_sheet=0,
                    source_path=source,
                    file_type="docx",
                    block_type="prose",
                    bbox=None,
                ))

        for table in doc.tables:
            all_blocks.extend(self._read_table(table, source))

        return all_blocks

    def _read_table(self, table: object, source_path: str) -> list[ExtractedBlock]:
        """Emit table_header and table_cell blocks for a python-docx Table."""
        blocks: list[ExtractedBlock] = []
        table_id = str(uuid.uuid4())
        headers: list[str] = []

        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if row_index == 0:
                    # First row is always headers — emit even if empty (preserves column index)
                    headers.append(text)
                    blocks.append(ExtractedBlock(
                        text=text,
                        page_or_sheet=0,
                        source_path=source_path,
                        file_type="docx",
                        block_type="table_header",
                        bbox=None,
                        table_id=table_id,
                        col_header=text,
                        row=row_index,
                        column=col_index,
                        row_index=row_index,
                    ))
                else:
                    if not text:
                        continue
                    col_header = headers[col_index] if col_index < len(headers) else None
                    blocks.append(ExtractedBlock(
                        text=text,
                        page_or_sheet=0,
                        source_path=source_path,
                        file_type="docx",
                        block_type="table_cell",
                        bbox=None,
                        table_id=table_id,
                        col_header=col_header,
                        row=row_index,
                        column=col_index,
                        row_index=row_index,
                    ))

        return blocks
